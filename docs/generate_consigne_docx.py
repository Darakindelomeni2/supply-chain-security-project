#!/usr/bin/env python3
"""
Génère la CONSIGNE ÉTUDIANTS (Word) du projet
"Sécurité de la chaîne d'approvisionnement logicielle (SLSA)".

Document autonome à envoyer aux étudiants : ils le lisent, comprennent le projet,
et savent exactement par où commencer.

Usage :
    pip install python-docx
    python generate_consigne_docx.py
    # -> Consigne-etudiants-Supply-Chain-Security.docx  (ignoré par git)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --------------------------------------------------------------------- palette / consts
NAVY  = RGBColor(0x0B, 0x1F, 0x3A)
TEAL  = RGBColor(0x0B, 0x81, 0x5A)
TEALB = "10B981"
RED   = RGBColor(0xC0, 0x2B, 0x2B)
GREY  = RGBColor(0x6B, 0x72, 0x80)
INK   = RGBColor(0x1F, 0x29, 0x37)
NAVYH = "0B1F3A"
LIGHT = "EEF2F6"
GREEN = "E7F7F0"
REDBG = "FBE9E9"
FONT  = "Calibri"

REPO = "https://github.com/aubinaso/supply-chain-security-project"

doc = Document()

# marges
for sec in doc.sections:
    sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)

# police de base
normal = doc.styles["Normal"]
normal.font.name = FONT
normal.font.size = Pt(11)
normal.font.color.rgb = INK

# recolorer les styles de titres
for lvl, col, sz in [("Title", NAVY, 26), ("Heading 1", NAVY, 16),
                     ("Heading 2", TEAL, 13), ("Heading 3", NAVY, 12)]:
    st = doc.styles[lvl]
    st.font.name = FONT
    st.font.color.rgb = col
    st.font.size = Pt(sz)
    st.font.bold = True


# --------------------------------------------------------------------- helpers
def _shade(element, fill_hex):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    element.append(shd)


def shade_cell(cell, fill_hex):
    _shade(cell._tc.get_or_add_tcPr(), fill_hex)


def para(text="", size=11, color=INK, bold=False, italic=False,
         align=None, space_after=6, space_before=0):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if text:
        r = p.add_run(text)
        r.font.size = Pt(size); r.font.color.rgb = color
        r.font.bold = bold; r.font.italic = italic; r.font.name = FONT
    return p


def runs(p, parts):
    """parts = liste de (texte, color, bold) ; ajoute des runs à un paragraphe."""
    for t, col, bold in parts:
        r = p.add_run(t)
        r.font.color.rgb = col; r.font.bold = bold
        r.font.size = Pt(11); r.font.name = FONT
    return p


def bullet(text, parts=None, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.space_after = Pt(3)
    if parts:
        runs(p, parts)
    else:
        r = p.add_run(text); r.font.size = Pt(11); r.font.name = FONT; r.font.color.rgb = INK
    return p


def banner(text_lines, fill_hex, text_color=RGBColor(0xFF, 0xFF, 0xFF)):
    """Bandeau coloré pleine largeur (table 1x1 shadée)."""
    t = doc.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0, 0)
    shade_cell(cell, fill_hex)
    cell.paragraphs[0].text = ""
    for i, (txt, sz, bold) in enumerate(text_lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(txt)
        r.font.size = Pt(sz); r.font.bold = bold; r.font.color.rgb = text_color; r.font.name = FONT
    _set_cell_margins(cell)
    return t


def callout(title, body_lines, fill_hex, bar_color=None, title_color=NAVY):
    """Encadré (astuce / attention) : table 1x1 shadée avec titre + corps."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade_cell(cell, fill_hex)
    _set_cell_margins(cell)
    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(3)
    r = p0.add_run(title); r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = title_color; r.font.name = FONT
    for line in body_lines:
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(line); r.font.size = Pt(10.5); r.font.color.rgb = INK; r.font.name = FONT
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def _set_cell_margins(cell, m=100):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(m)); el.set(qn("w:type"), "dxa")
        tcMar.append(el)
    tcPr.append(tcMar)


def table(headers, rows_data, widths=None, header_fill=NAVYH,
          header_color=RGBColor(0xFF, 0xFF, 0xFF), zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # en-tête
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        shade_cell(hdr[i], header_fill)
        _set_cell_margins(hdr[i], 80)
        p = hdr[i].paragraphs[0]
        r = p.add_run(h); r.font.bold = True; r.font.size = Pt(10.5)
        r.font.color.rgb = header_color; r.font.name = FONT
    # lignes
    for ri, row in enumerate(rows_data):
        cells = t.add_row().cells
        for ci, val in enumerate(row):
            _set_cell_margins(cells[ci], 80)
            if zebra and ri % 2 == 1:
                shade_cell(cells[ci], LIGHT)
            p = cells[ci].paragraphs[0]
            # val peut être (texte, bold, color)
            if isinstance(val, tuple):
                txt, bold, col = val
            else:
                txt, bold, col = val, False, INK
            r = p.add_run(txt); r.font.size = Pt(10); r.font.bold = bold
            r.font.color.rgb = col; r.font.name = FONT
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    return t


def mono(text_lines):
    """Bloc « code » : paragraphe grisé en Consolas."""
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade_cell(cell, "0B1F3A")
    _set_cell_margins(cell, 120)
    for i, line in enumerate(text_lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xE6, 0xF0, 0xEA)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def spacer(pts=4):
    doc.add_paragraph().paragraph_format.space_after = Pt(pts)


# ===================================================================== EN-TÊTE / TITRE
banner([
    ("PROJET TECHNIQUE — 5ᵉ ANNÉE DEVOPS / CLOUD / INFRASTRUCTURE", 11, True),
    ("Sécuriser la chaîne d'approvisionnement logicielle (SLSA)", 20, True),
    ("Consigne du projet — à lire avant de commencer", 12, False),
], NAVYH)
spacer(2)
p = para("", space_after=8)
runs(p, [("Durée : ", NAVY, True), ("3 jours ", INK, False),
         ("(~1,5 jour de réalisation + QCM + soutenance)   ·   ", GREY, False),
         ("Travail en groupe : ", NAVY, True), ("2 à 4 étudiants   ·   ", INK, False),
         ("QCM : ", NAVY, True), ("individuel", INK, False)])

callout("En une phrase",
        ["Vous avez déjà construit des pipelines CI/CD. Ici, vous répondez à la question que se "
         "posent toutes les équipes DevSecOps : « comment PROUVER que l'image qui tourne en "
         "production est bien celle que NOUS avons construite — et pas une version piégée ? »"],
        GREEN, title_color=TEAL)

# ===================================================================== 1. CONTEXTE
doc.add_heading("1. Contexte — pourquoi ce projet", level=1)
para("Vous savez tous faire : build → test → scan → deploy. Mais une fois l'image en production, "
     "qu'est-ce qui garantit qu'elle n'a pas été altérée entre le build et le déploiement ? "
     "Un « docker pull » ne vérifie rien, et « le scan était vert » ne prouve pas que l'image "
     "DÉPLOYÉE est celle qui a été scannée.")
para("Les attaques récentes ne visent plus votre application : elles visent votre chaîne de "
     "fabrication logicielle.", bold=True, color=NAVY, space_after=4)
bullet("", [("SolarWinds (2020) — ", NAVY, True),
            ("code malveillant injecté dans le build, signé par l'éditeur, poussé à 18 000 clients.", INK, False)])
bullet("", [("Codecov (2021) — ", NAVY, True),
            ("script CI modifié exfiltrant les secrets des pipelines de milliers de projets.", INK, False)])
bullet("", [("Dependency confusion (2021) — ", NAVY, True),
            ("faux paquets « internes » publiés sur les registries publics.", INK, False)])
bullet("", [("XZ Utils / liblzma (2024) — ", NAVY, True),
            ("backdoor introduite sur 3 ans dans une dépendance open source très répandue.", INK, False)])
para("La réponse de l'industrie — que vous allez mettre en œuvre — s'appuie sur : SBOM, "
     "signature (Sigstore/cosign), attestations de provenance (SLSA) et contrôle à l'admission "
     "(policy-as-code).", space_before=4)

# ===================================================================== 2. OBJECTIF
doc.add_heading("2. Objectif — ce que vous devez livrer (POC)", level=1)
para("Sur l'application fournie, livrez un POC opérationnel démontrant :", space_after=4)
for i, txt in enumerate([
    "Un SBOM généré et un scan de vulnérabilités qui CASSE le build en cas de CVE critique.",
    "L'image SIGNÉE (cosign) avec deux attestations attachées : le SBOM et la provenance (SLSA).",
    "Un cluster Kubernetes (kind/k3s) avec Kyverno qui EXIGE signature + attestations + registry "
    "autorisé + interdiction du tag :latest.",
    "Une démonstration attaque/défense : une image non signée OU modifiée après signature est "
    "REJETÉE par le cluster, capture à l'appui.",
    "Un threat model court de la chaîne + l'argument « quel contrôle mitige quelle attaque ».",
]):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(txt); r.font.size = Pt(11); r.font.name = FONT; r.font.color.rgb = INK

# ===================================================================== 3. LA CHAÎNE
doc.add_heading("3. La chaîne que vous allez construire", level=1)
mono([
    " code ─► build ─► SBOM (Syft) ─► scan (Grype) ─► SIGNATURE (cosign)",
    "                                                     │",
    "                                                     ├─► attestation SBOM",
    "                                                     └─► attestation de PROVENANCE (SLSA)",
    "                                                            │  push ─► GHCR (registry)",
    "   ┌────────────────────────────────────────────────────────┘",
    "   ▼",
    " Cluster Kubernetes (kind/k3s) + KYVERNO (admission control)",
    "   ├─ signée par NOTRE identité ?           sinon ─► ❌ REFUSÉE",
    "   ├─ attestation de provenance présente ?   sinon ─► ❌ REFUSÉE",
    "   ├─ registry autorisé + par digest ?       sinon ─► ❌ REFUSÉE",
    "   └─ pas de CVE critique / pas de :latest ?  sinon ─► ❌ REFUSÉE",
])
p = para("", space_after=8)
runs(p, [("Résultat : ", NAVY, True),
         ("image légitime → ✅ le pod tourne ; image non signée ou modifiée → ❌ le cluster la bloque, en direct.", INK, False)])

doc.add_heading("Les 4 briques (le vrai sujet)", level=2)
table(
    ["Brique", "Outil", "Ce qu'elle apporte"],
    [
        [("1 · SBOM", True, NAVY), "Syft + Grype", "L'inventaire exact des composants de l'image ; « suis-je affecté par la CVE du jour ? » en secondes."],
        [("2 · Signature", True, NAVY), "cosign / Sigstore", "Preuve cryptographique « c'est bien nous ». Mode keyless (identité OIDC), journalisé dans Rekor."],
        [("3 · Attestations", True, NAVY), "cosign attest", "Affirmations signées attachées à l'image : le SBOM + la provenance (qui/quoi/d'où/quand)."],
        [("4 · Admission", True, NAVY), "Kyverno", "Le gardien du cluster : vérifie signature + attestations et REFUSE l'inconnu."],
    ],
    widths=[3.2, 3.3, 10.5],
)
spacer(2)
callout("À retenir",
        ["Le cluster ne fait JAMAIS confiance à un tag. Il exige une preuve cryptographique liée au "
         "digest de l'image. Si un octet change après signature, le digest change, la signature ne "
         "correspond plus, et Kyverno refuse. C'est le passage de « on scanne » à « on VÉRIFIE et on BLOQUE »."],
        LIGHT, title_color=NAVY)

# ===================================================================== 4. PLANNING
doc.add_heading("4. Déroulé des 3 jours", level=1)
table(
    ["Quand", "Thème", "Activité"],
    [
        [("Jour 1", True, NAVY), ("Chaîne vérifiable", False, TEAL), "Labs 0→2 : build, SBOM, scan, signature, attestations."],
        [("Jour 2 matin", True, NAVY), ("Le cluster qui refuse", False, TEAL), "Labs 3→4 : Kyverno + attaque/défense.  ← fin du temps de projet (1,5 j)"],
        [("Jour 2 aprèm", True, NAVY), ("Intégration & rédaction", False, TEAL), "Lab 5 (CI GitHub Actions, bonus) + rapport + threat model."],
        [("Jour 3 matin", True, NAVY), ("QCM individuel", False, TEAL), "Répétition de la démo + QCM (25-30 min)."],
        [("Jour 3 aprèm", True, NAVY), ("Soutenances", False, RED), "12 min démo/présentation + 5 min de questions par groupe."],
    ],
    widths=[2.8, 4.2, 10.0],
)

# ===================================================================== 5. LES LABS
doc.add_heading("5. Votre parcours : les 5 labs guidés", level=1)
para("Le cœur des 1,5 jour. Chaque lab a des critères de sortie clairs. Ils sont dans le dépôt, "
     "dossier labs/.", space_after=4)
table(
    ["Lab", "Titre", "Ce que vous faites"],
    [
        [("Lab 0", True, NAVY), "Setup & première image", "Outils, fork, build, push sur GHCR, notion de digest."],
        [("Lab 1", True, NAVY), "SBOM & scan qui casse", "Syft (SBOM) + Grype (gate qui stoppe sur CVE critique)."],
        [("Lab 2", True, NAVY), "Signer & attester", "cosign : signature + attestations SBOM et provenance."],
        [("Lab 3", True, NAVY), "Le cluster qui refuse", "kind + Kyverno : exiger signature, provenance, registry, pas de :latest."],
        [("Lab 4", True, NAVY), "Attaque / défense", "Image non signée / modifiée → BLOQUÉE. Captures pour la démo."],
        [("Lab 5", True, NAVY), "CI de bout en bout (bonus)", "Tout automatiser en GitHub Actions (keyless) → vers SLSA L2."],
    ],
    widths=[2.0, 4.8, 10.2],
)

# ===================================================================== 6. PRÉREQUIS / CODE FOURNI
doc.add_heading("6. Prérequis & code fourni", level=1)
para("Vous ne partez PAS d'une page blanche. Vous forkez le dépôt du projet, qui contient déjà "
     "l'application, les politiques, les manifestes et le pipeline de référence. Votre travail : "
     "faire fonctionner la chaîne, la compléter, et la personnaliser (remplacer les « <votre-user> »).",
     space_after=6)
table(
    ["Fourni (à ne pas réécrire)", "À produire / personnaliser par vous"],
    [
        ["app/ — API Flask + Dockerfile + tests", "SBOM, scan, signature, attestations (labs 1-2)"],
        ["policies/kyverno/ — 4 politiques prêtes", "Y coller votre cosign.pub / identité + <votre-user>"],
        ["k8s/deployment.yaml", "Y mettre votre image par digest"],
        [".github/workflows/ — pipeline de référence", "L'activer sur votre fork (lab 5)"],
        ["labs/, docs/, evaluation/, livrables/", "Remplir les livrables (rapport, threat model)"],
    ],
    widths=[8.5, 8.5],
)
spacer(2)
para("Outils à installer (tout tourne en local, aucun cloud requis) :", bold=True, color=NAVY, space_after=3)
bullet("Docker · kind (ou k3s) · kubectl")
bullet("Syft · Grype · cosign")
bullet("git · jq · un compte GitHub (+ un token PAT avec le scope write:packages pour le registry GHCR)")
para("Détail des versions et installation par OS : docs/01-prerequis-setup.md dans le dépôt.",
     italic=True, color=GREY, size=10)

# ===================================================================== 7. LIVRABLES & BARÈME
doc.add_heading("7. Livrables attendus & barème", level=1)
table(
    ["#", "Livrable", "Format"],
    [
        [("L1", True, NAVY), "POC fonctionnel (dépôt forké : app, SBOM, signature, attestations, politiques, manifs)", "Repo GitHub"],
        [("L2", True, NAVY), "Rapport court (5-8 p.) : ce que vous avez fait, comment le vérifier, niveau SLSA, limites", "PDF / Markdown"],
        [("L3", True, NAVY), "Threat model : attaques → contrôles → couverture", "PDF / Markdown"],
        [("L4", True, NAVY), "Démo attaque/défense (captures ou vidéo : une image rejetée + une acceptée)", "dans le rapport"],
        [("L5", True, NAVY), "Soutenance : présentation + démo live", "12 min + 5 min Q/R"],
        [("QCM", True, NAVY), "QCM individuel sur les concepts", "sur place (J3)"],
    ],
    widths=[1.6, 12.4, 3.0],
)
spacer(2)
para("Barème (100 %) :", bold=True, color=NAVY, space_after=3)
table(
    ["Composante", "Poids"],
    [
        ["POC & démo (L1 + L4 + démo de soutenance)", ("35 %", True, NAVY)],
        ["Rapport + threat model (L2 + L3)", ("25 %", True, NAVY)],
        ["Soutenance (L5)", ("20 %", True, NAVY)],
        ["QCM individuel", ("20 %", True, NAVY)],
    ],
    widths=[13.0, 4.0],
)

# ===================================================================== 8. CRITÈRES DE RÉUSSITE
doc.add_heading("8. Critères de réussite du POC (auto-évaluation)", level=1)
para("Cochez avant la soutenance — c'est aussi ce que l'encadrant regardera :", space_after=4)
for txt in [
    "Un SBOM (SPDX ou CycloneDX) est généré pour l'image.",
    "Le scan CASSE le build en présence d'une CVE critique corrigeable.",
    "L'image est signée et « cosign verify » réussit avec votre identité.",
    "Une attestation SBOM et une attestation de provenance sont attachées et vérifiables.",
    "Le cluster ACCEPTE votre image signée et conforme.",
    "Le cluster REFUSE une image non signée (message d'erreur Kyverno à l'appui).",
    "Le cluster REFUSE une image modifiée après signature (le digest ne correspond plus).",
    "Le cluster REFUSE le tag :latest et/ou un registry non autorisé.",
    "Tout est reproductible : « kind create » + « kubectl apply » reconstruit la démo.",
]:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("☐  " + txt); r.font.size = Pt(10.5); r.font.name = FONT; r.font.color.rgb = INK

# ===================================================================== 9. RÈGLES & CONSEILS
doc.add_heading("9. Règles & conseils", level=1)
bullet("", [("Groupes de 2 à 4. ", NAVY, True),
            ("Chaque membre doit commiter (la traçabilité Git, c'est justement le sujet du projet).", INK, False)])
bullet("", [("Ne commitez JAMAIS de secret. ", RED, True),
            ("La clé de signature cosign.key est déjà dans le .gitignore — laissez-la ignorée.", INK, False)])
bullet("", [("Travaillez par digest, pas par tag. ", NAVY, True),
            ("Signer/déployer « @sha256:… » est le cœur de la garantie d'intégrité.", INK, False)])
bullet("", [("Plan B pour la démo : ", NAVY, True),
            ("enregistrez une capture vidéo de votre séquence attaque/défense en fin de J2. Si le "
             "live plante en soutenance, vous avez une preuve.", INK, False)])
bullet("", [("Soyez honnêtes sur SLSA : ", NAVY, True),
            ("annoncez le niveau réellement atteint (L1/L2) et ce qui reste contournable.", INK, False)])

# ===================================================================== 10. PAR OÙ COMMENCER
doc.add_heading("10. Par où commencer (maintenant)", level=1)
for i, txt in enumerate([
    "Constituez votre groupe et forkez le dépôt du projet dans votre compte GitHub.",
    "Clonez votre fork et installez les outils (docs/01-prerequis-setup.md).",
    "Lisez docs/00-presentation-projet.md pour la vue d'ensemble.",
    "Enchaînez les labs : labs/lab0-setup.md → lab4 (puis lab5 en bonus).",
    "Préparez au fil de l'eau vos livrables (templates dans livrables/).",
]):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(txt); r.font.size = Pt(11); r.font.name = FONT; r.font.color.rgb = INK

spacer(2)
p = para("", space_after=4)
runs(p, [("Dépôt du projet : ", NAVY, True), (REPO, TEAL, True)])
callout("Objectif de fin de Jour 1",
        ["Une image SIGNÉE, vérifiable (SBOM + provenance attachés), poussée dans votre registry. "
         "Si vous y êtes, la partie « cluster qui refuse » du Jour 2 s'enchaîne naturellement."],
        GREEN, title_color=TEAL)

para("On ne fait pas confiance — on vérifie.", bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER,
     space_before=6, size=13)

# --------------------------------------------------------------------- save
out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                   "Consigne-etudiants-Supply-Chain-Security.docx")
doc.save(out)
print("OK ->", out)
